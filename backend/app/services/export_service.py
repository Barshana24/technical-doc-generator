import io
import re
from typing import Optional
from app.models.documentation import GeneratedDocument, Project
from app.utils.logger import logger


def _strip_markdown(text: str) -> str:
    """Convert Markdown to plain text for PDF/DOCX."""
    text = re.sub(r"#{1,6}\s+", "", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`{1,3}[^`]*`{1,3}", lambda m: m.group(0).replace("`", ""), text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    return text


def generate_pdf(project: Project, documents: list[GeneratedDocument]) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Preformatted,
            HRFlowable, PageBreak,
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError as exc:
        raise RuntimeError("reportlab is not installed") from exc

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title=f"{project.name} — Technical Documentation",
        author="Technical Documentation Generator",
    )

    base_styles = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "DocTitle", parent=base_styles["Title"],
            fontSize=24, spaceAfter=12, textColor=colors.HexColor("#1a1a2e"),
        ),
        "h1": ParagraphStyle(
            "Heading1", parent=base_styles["Heading1"],
            fontSize=18, spaceAfter=8, textColor=colors.HexColor("#16213e"),
        ),
        "h2": ParagraphStyle(
            "Heading2", parent=base_styles["Heading2"],
            fontSize=14, spaceAfter=6, textColor=colors.HexColor("#0f3460"),
        ),
        "body": ParagraphStyle(
            "Body", parent=base_styles["Normal"],
            fontSize=10, leading=14, spaceAfter=6,
        ),
        "code": ParagraphStyle(
            "Code", parent=base_styles["Code"],
            fontSize=8, leading=10, fontName="Courier",
            backColor=colors.HexColor("#f4f4f4"),
            leftIndent=10, rightIndent=10,
        ),
        "meta": ParagraphStyle(
            "Meta", parent=base_styles["Normal"],
            fontSize=9, textColor=colors.grey, spaceAfter=4,
        ),
    }

    story = []

    # Cover page
    story.append(Spacer(1, 1.5 * inch))
    story.append(Paragraph(project.name, styles["title"]))
    story.append(Paragraph("Technical Documentation", styles["h1"]))
    story.append(Spacer(1, 0.3 * inch))

    if project.language_stats:
        langs = ", ".join(f"{k} ({v})" for k, v in list(project.language_stats.items())[:5])
        story.append(Paragraph(f"Languages: {langs}", styles["meta"]))
    if project.file_count:
        story.append(Paragraph(f"Files analyzed: {project.file_count}", styles["meta"]))

    story.append(Spacer(1, 0.2 * inch))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(PageBreak())

    # Documents
    for gd in documents:
        story.append(Paragraph(gd.title, styles["h1"]))
        story.append(Paragraph(f"Type: {gd.doc_type.value}", styles["meta"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#eeeeee")))
        story.append(Spacer(1, 0.15 * inch))

        content = gd.content or ""
        in_code_block = False
        code_lines = []

        for line in content.splitlines():
            if line.startswith("```"):
                if in_code_block:
                    # flush code block
                    story.append(Preformatted("\n".join(code_lines), styles["code"]))
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if line.startswith("# "):
                story.append(Paragraph(line[2:], styles["h1"]))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], styles["h2"]))
            elif line.startswith("### "):
                story.append(Paragraph(line[4:], styles["h2"]))
            elif line.strip():
                safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe, styles["body"]))
            else:
                story.append(Spacer(1, 0.1 * inch))

        if in_code_block and code_lines:
            story.append(Preformatted("\n".join(code_lines), styles["code"]))

        if gd.mermaid_diagram:
            story.append(Spacer(1, 0.1 * inch))
            story.append(Paragraph("UML Diagram (Mermaid Source)", styles["h2"]))
            story.append(Preformatted(gd.mermaid_diagram, styles["code"]))

        story.append(PageBreak())

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def generate_docx(project: Project, documents: list[GeneratedDocument]) -> bytes:
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc

    docx = DocxDocument()

    # Styles
    docx.core_properties.title = f"{project.name} — Technical Documentation"
    docx.core_properties.author = "Technical Documentation Generator"

    # Cover
    cover = docx.add_heading(project.name, level=0)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = docx.add_paragraph("Technical Documentation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    if project.language_stats:
        langs = ", ".join(f"{k} ({v})" for k, v in list(project.language_stats.items())[:5])
        p = docx.add_paragraph(f"Languages: {langs}")
        p.runs[0].font.size = Pt(10)
    if project.file_count:
        p = docx.add_paragraph(f"Files analyzed: {project.file_count}")
        p.runs[0].font.size = Pt(10)

    docx.add_page_break()

    for gd in documents:
        docx.add_heading(gd.title, level=1)
        meta_p = docx.add_paragraph(f"Type: {gd.doc_type.value}")
        meta_p.runs[0].font.size = Pt(9)
        meta_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        content = gd.content or ""
        in_code_block = False
        code_lines = []

        for line in content.splitlines():
            if line.startswith("```"):
                if in_code_block:
                    code_text = "\n".join(code_lines)
                    p = docx.add_paragraph(style="No Spacing")
                    run = p.add_run(code_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if line.startswith("# "):
                docx.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                docx.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                docx.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                p = docx.add_paragraph(style="List Bullet")
                p.add_run(line[2:])
            elif line.strip():
                docx.add_paragraph(line)

        if in_code_block and code_lines:
            p = docx.add_paragraph(style="No Spacing")
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(8)

        if gd.mermaid_diagram:
            docx.add_heading("UML Diagram (Mermaid Source)", level=2)
            p = docx.add_paragraph(style="No Spacing")
            run = p.add_run(gd.mermaid_diagram)
            run.font.name = "Courier New"
            run.font.size = Pt(8)

        docx.add_page_break()

    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    return buffer.read()
